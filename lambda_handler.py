import os
import json
import time
import logging
import requests
import ast
import pandas as pd
import numpy as np
from datetime import datetime
import subprocess  # <-- For running LibreOffice headless


# ----------------------------------------------------------------------------
# HELPER – Due-date parsing
# ----------------------------------------------------------------------------

def safe_parse_due(due_val):
    """Parse ISO string or epoch seconds into YYYY-MM-DD; return blank on failure."""
    if not due_val or (isinstance(due_val, float) and pd.isna(due_val)):
        return ""
    try:
        return str(datetime.fromisoformat(str(due_val).replace("Z", "")).date())
    except Exception:
        pass
    try:
        return str(datetime.fromtimestamp(float(due_val)).date())
    except Exception:
        return str(due_val)
import boto3
from botocore.exceptions import ClientError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ----------------------------------------------------------------------------
# CONFIG
# ----------------------------------------------------------------------------

SECRET_NAME = "ProjectPlaceAPICredentials"
REGION = "us-east-2"
PROJECTPLACE_API_URL = "https://api.projectplace.com"

# Files & Paths
OUTPUT_EXCEL = "/tmp/Acta_de_Seguimiento.xlsx"

# Environment Variables
DYNAMO_TABLE = os.getenv("DYNAMODB_TABLE_NAME", "ProjectPlace_DataExtrator_landing_table_v3")
S3_BUCKET = os.getenv("S3_BUCKET_NAME", "projectplace-dv-2025-x9a7b")

# BRAND_COLOR_HEADER changed from #2E86C1 → #4AC795
BRAND_COLOR_HEADER = "4AC795"
LIGHT_SHADE_2X2 = "FAFAFA"

# Location for your logo image. Make sure it's included in your Docker image.
LOGO_IMAGE_PATH = os.path.join(os.path.dirname(__file__), "logo", "company_logo.png")


def lambda_handler(event, context):
    """
    1) Load secrets -> get token
    2) Fetch *all active enterprise projects* dynamically
    3) Generate Excel with tasks (cards) for each project + comments
    4) snippet_filter => col_id=1, parse 'creator' => 'creator_name', parse last comment
    5) Multi-doc creation (one doc per project)
    6) Leader name => from 'creator_name'
    7) Relaxed date parse for FECHA => if parse fails, display raw comment
    8) Reordered columns for COMPROMISOS => ["COMPROMISO","RESPONSABLE","FECHA"]
    9) Store in Dynamo + Upload to S3
    10) (OPTIONAL) Convert docx to PDF + upload.
    """
    logger.info("🚀 Starting final Acta generation with dynamic project list + leader logic.")

    secret_dict = load_secrets()
    client_id = secret_dict.get("PROJECTPLACE_ROBOT_CLIENT_ID", "")
    client_secret = secret_dict.get("PROJECTPLACE_ROBOT_CLIENT_SECRET", "")
    if not client_id or not client_secret:
        return {"statusCode": 500, "body": "Missing ProjectPlace credentials."}

    # 1) Get Robot Token
    token = get_robot_access_token(client_id, client_secret)
    if not token:
        return {"statusCode": 500, "body": "Failed to get ProjectPlace token."}

    # 2) Dynamically fetch *active* enterprise projects
    projects = get_all_account_projects(token, include_archived=False)
    if not projects:
        msg = "No active projects returned. Possibly no new or user has limited visibility."
        logger.warning(msg)
        return {"statusCode": 200, "body": msg}

    # 3) Generate Excel with tasks (cards) for each project + comments
    excel_path = generate_excel_report(token, projects)
    if not excel_path:
        return {"statusCode": 500, "body": "No Excel generated (no cards?)."}

    # 4) Build DataFrame from the Excel
    df = pd.read_excel(excel_path)
    
    # 🔍  see what came back from Excel
    logger.info(f"[DEBUG] Excel load → rows={len(df)}, cols={list(df.columns)}")
    
    if df.empty:
        logger.warning("Excel is empty—no tasks to process.")
        return {"statusCode": 200, "body": "No tasks found in Excel."}
    
    # 5) Store in Dynamo
    store_in_dynamodb(df)
    
    # 6) snippet => filter + parse
    df = snippet_filter(df)
    logger.info(f"[DEBUG] After snippet_filter → rows={len(df)}")   # <— new breadcrumb
    
    if df.empty:
        logger.warning("No tasks remain after snippet filter.")
        return {"statusCode": 200, "body": "No tasks remain after snippet filter."}
    
    # 7) Multi-doc creation
    grouped = df.groupby("project_id")
    doc_count = 0
    for pid, project_df in grouped:
        logger.info(f"[DEBUG] Build acta pid={pid} rows={len(project_df)}")  # <— new breadcrumb
        doc_path = build_acta_for_project(pid, project_df)
        if not doc_path:
            logger.error(f"❌ build_acta returned None for {pid}")
            continue
        doc_count += 1
        
        # Upload doc
        first_row = project_df.iloc[0]
        p_name = str(first_row.get("project_name", "UnknownProject"))
        safe_proj = p_name.replace("/", "_").replace(" ", "_")
        s3_key_docx = f"actas/Acta_{safe_proj}_{pid}.docx"
        upload_file_to_s3(doc_path, s3_key_docx)

        # (OPTIONAL) Convert to PDF and upload
        try:
            pdf_path = convert_docx_to_pdf(doc_path)
            if pdf_path:
                # e.g. "actas/Acta_ProjectName_1234.pdf"
                s3_key_pdf = f"actas/Acta_{safe_proj}_{pid}.pdf"
                upload_file_to_s3(pdf_path, s3_key_pdf)
        except Exception as exc:
            logger.error(f"PDF conversion failed for doc {doc_path}: {exc}")
            
    # 8) Upload Excel as well
    s3_excel_key = "actas/Acta_de_Seguimiento.xlsx"
    upload_file_to_s3(excel_path, s3_excel_key)

    msg = {
        "message": f"Multi-Acta creation done. Docs created={doc_count}",
        "excel_s3": f"s3://{S3_BUCKET}/{s3_excel_key}"
    }
    return {"statusCode": 200, "body": json.dumps(msg)}


# ----------------------------------------------------------------------------
# 1) SECRETS & OAUTH
# ----------------------------------------------------------------------------
def load_secrets():
    sm = boto3.client("secretsmanager", region_name=REGION)
    try:
        resp = sm.get_secret_value(SecretId=SECRET_NAME)
        return json.loads(resp["SecretString"])
    except ClientError as e:
        logger.error(f"Secrets error: {str(e)}")
        return {}

def get_robot_access_token(client_id, client_secret):
    url = f"{PROJECTPLACE_API_URL}/oauth2/access_token"
    data = {
        "grant_type": "client_credentials",
        "client_id": client_id,
        "client_secret": client_secret
    }
    try:
        resp = requests.post(url, data=data)
        resp.raise_for_status()
        return resp.json().get("access_token")
    except requests.exceptions.RequestException as e:
        logger.error(f"Token fetch error: {str(e)}")
        return None

# ----------------------------------------------------------------------------
# 2) ENTERPRISE PROJECTS
# ----------------------------------------------------------------------------
def get_all_account_projects(token, include_archived=False):
    """
    Calls /1/account/projects to list projects for the entire enterprise.
    If include_archived=False, skip archived.
    """
    url = f"{PROJECTPLACE_API_URL}/1/account/projects"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    params = {}
    if include_archived:
        params["include_archived"] = 1

    try:
        r = requests.get(url, headers=headers, params=params)
        r.raise_for_status()
        data = r.json()
        if not isinstance(data, dict):
            logger.warning(f"Projects response not dict: {type(data)} => {data}")
            return []

        projects_list = data.get("projects", [])
        if not projects_list:
            logger.info("No projects found in /1/account/projects.")
            return []

        if not include_archived:
            active_projects = [p for p in projects_list if not p.get("archived", False)]
            logger.info(f"Fetched {len(active_projects)} active projects.")
            return active_projects
        else:
            logger.info(f"Fetched {len(projects_list)} total projects (including archived).")
            return projects_list

    except requests.exceptions.RequestException as e:
        logger.error(f"Error listing enterprise projects: {e}")
        return []

# ----------------------------------------------------------------------------
# 3) EXCEL GENERATION
# ----------------------------------------------------------------------------
def generate_excel_report(token, projects):
    if not projects:
        logger.warning("No projects provided to generate_excel_report.")
        return None

    all_cards = []
    headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}

    for project_info in projects:
        pid = str(project_info.get("id"))
        p_name = project_info.get("name", "Unnamed Project")
        cards_url = f"{PROJECTPLACE_API_URL}/1/projects/{pid}/cards"

        try:
            resp = requests.get(cards_url, headers=headers)
            resp.raise_for_status()
            cards = resp.json()
            for c in cards:
                row = dict(c)
                cid = c.get("id")

                cmts = fetch_comments_for_card(token, cid)
                # Capture label_id
                label_val = None
                if "label_id" in c:
                    label_val = c.get("label_id")
                elif isinstance(c.get("labels"), list) and c["labels"]:
                    label_val = c["labels"][0].get("id")
                row["label_id"] = label_val
                row["Comments"] = str(cmts)

                row["project_id"] = pid
                row["project_name"] = p_name
                row["archived"] = project_info.get("archived", False)

                all_cards.append(row)

            logger.info(f"Fetched {len(cards)} cards from project '{p_name}' ({pid}).")

        except Exception as e:
            logger.error(f"Error fetching cards for project {pid}: {str(e)}")

    if not all_cards:
        logger.warning("No cards found across all projects.")
        return None

    df = pd.DataFrame(all_cards)
    df.to_excel(OUTPUT_EXCEL, index=False)
    logger.info(f"✅ Wrote {len(df)} rows to {OUTPUT_EXCEL}")
    return OUTPUT_EXCEL


def fetch_comments_for_card(token, card_id):
    if not card_id:
        return []
    url = f"{PROJECTPLACE_API_URL}/1/cards/{card_id}/comments"
    headers = {"Authorization": f"Bearer {token}", "Accept":"application/json"}
    out = []
    try:
        r = requests.get(url, headers=headers)
        r.raise_for_status()
        for c in r.json():
            out.append(c.get("text","N/A"))
    except Exception as e:
        logger.error(f"Failed to fetch comments for card {card_id}: {str(e)}")
    return out

# ----------------------------------------------------------------------------
# 4) DYNAMO
# ----------------------------------------------------------------------------
def store_in_dynamodb(df):
    if df.empty:
        return
    ddb = boto3.resource("dynamodb", region_name=REGION)
    table = ddb.Table(DYNAMO_TABLE)
    inserted = 0
    for idx, row in df.iterrows():
        pid = row.get("project_id")
        if not pid or pd.isna(pid):
            continue
        item = {
            "project_id": str(pid),
            "card_id": str(row.get("id","N/A")),
            "title": str(row.get("title","N/A")),
                    "label_id": str(row.get("label_id","")),
            "timestamp": int(time.time())
        }
        table.put_item(Item=item)
        inserted += 1
    logger.info(f"Inserted {inserted} items into {DYNAMO_TABLE}")


# ----------------------------------------------------------------------------
# 5) SNIPPET FILTER
# ----------------------------------------------------------------------------
def snippet_filter(df):
    if "column_id" in df.columns:
        df = df[df["column_id"] == 1].copy()

    df["comments_parsed"] = df["Comments"].apply(parse_last_comment)

    if "project" in df.columns:
        df["project_dict"] = df["project"].apply(parse_dict_column)
        df["project_id"] = df["project_dict"].apply(lambda p: p.get("id", None))
        df["project_name"] = df["project_dict"].apply(lambda p: p.get("name","Unknown Project"))

    if "creator" in df.columns:
        df["creator_dict"] = df["creator"].apply(parse_dict_column)
        df["creator_name"] = df["creator_dict"].apply(lambda c: c.get("name","N/A"))
    else:
        df["creator_name"] = "N/A"

    if "planlet" in df.columns:
        df["planlet_dict"] = df["planlet"].apply(parse_dict_column)
        df["planlet_name"] = df["planlet_dict"].apply(lambda d: d.get("label", d.get("name","")))
        df["wbs_str"] = df["planlet_dict"].apply(lambda d: d.get("wbs_id",""))
        df["wbs_tuple"] = df["wbs_str"].apply(parse_wbs_id)
        df = df[df["wbs_tuple"].apply(len)>0].copy()

    df = df[df["comments_parsed"].str.strip() != ""].copy()
    return df

def parse_dict_column(raw_value):
    try:
        return ast.literal_eval(str(raw_value))
    except:
        return {}

def parse_last_comment(raw_value):
    try:
        arr = ast.literal_eval(str(raw_value))
        if isinstance(arr, list) and len(arr) > 0:
            return str(arr[-1])
    except:
        pass
    return ""

def parse_wbs_id(wbs_str):
    if not wbs_str:
        return ()
    parts = wbs_str.split(".")
    out = []
    for p in parts:
        try:
            out.append(int(p.strip().rstrip(",.")))
        except:
            pass
    return tuple(out)

# ----------------------------------------------------------------------------
# 6) BUILD ACTA => includes COMPROMISOS
# ----------------------------------------------------------------------------
def build_acta_for_project(pid, project_df):
    if project_df.empty:
        return None

    if "wbs_tuple" in project_df.columns:
        project_df = project_df.sort_values("wbs_tuple", ascending=True)

    first_row = project_df.iloc[0]
    project_name = str(first_row.get("project_name","Unknown Project"))
    leader_name = str(first_row.get("creator_name","N/A"))

    doc = Document()
    set_document_margins_and_orientation(doc)
    add_page_x_of_y_footer(doc)

    add_unified_visual_header(doc, "ACTA DE SEGUIMIENTO", LOGO_IMAGE_PATH)  # 👈 Adds legal fields as shown in client doc
    doc.add_paragraph()

    date_now = datetime.now().strftime("%m/%d/%Y")
    date_text = f"FECHA DEL ACTA: {date_now}"
    add_two_by_two_table(doc, date_text, project_name, str(pid), leader_name, shade_color=LIGHT_SHADE_2X2)

    doc.add_paragraph()
    add_asistencia_table(doc, project_df)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    add_section_header(doc, "ESTADO DEL PROYECTO Y COMENTARIOS")
    add_project_status_table(doc, project_df)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    add_section_header(doc, "COMPROMISOS")
    add_commitments_table(doc, project_df)
    doc.add_paragraph()

    safe_proj_name = project_name.replace(" ", "_").replace("/", "_")
    doc_path = f"/tmp/Acta_{safe_proj_name}_{pid}.docx"
    doc.save(doc_path)
    logger.info(f"Created doc => {doc_path}")
    return doc_path

def add_project_status_table(doc, df):
    # --- filter rows ---------------------------------------------------
    df = df[df["label_id"] != 0]
    df = df[df.get("board_name", "") != "COMPROMISOS"].copy()
    df = df[~df["planlet_name"].isin(
        ["ASISTENCIA", "ASISTENCIA CLIENTE", "ASISTENCIA IKUSI"]
    )].copy()

    # --- collect rows for the table -----------------------------------
    table_data = []
    for _, row in df.iterrows():
        hito        = str(row.get("planlet_name", ""))
        actividades = str(row.get("title", ""))
        desarrollo  = str(row.get("comments_parsed", ""))
        table_data.append([hito, actividades, desarrollo])

    # --- build the actual Word table (one-time) ------------------------
    table = doc.add_table(rows=1, cols=3)
    table.style  = "Table Grid"
    table.autofit = False

    hdr_row = table.rows[0]
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr_row.height      = Inches(0.45)

    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(2.5)
    table.columns[2].width = Inches(5.0)

    # CHANGED: "HITO" => "HITO/TEMA"
    headers = ["HITO/TEMA", "ACTIVIDADES", "DESARROLLO"]  # <--- CHANGED
    hdr_cells = hdr_row.cells
    for i, hdr_text in enumerate(headers):
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].text = hdr_text
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Verdana"
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), BRAND_COLOR_HEADER)
        hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

    for row_idx, row_data in enumerate(table_data):
        new_cells = table.add_row().cells
        for col_idx, val in enumerate(row_data):
            new_cells[col_idx].text = val
            p = new_cells[col_idx].paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(10)
            run.font.name = "Verdana"


def add_commitments_table(doc: Document, df: pd.DataFrame) -> None:
    """
    Build the COMPROMISOS table.

    • New logic:   label_id == 0  AND column_id == 1
    • Legacy:      board_name == "COMPROMISOS"     (column_id may be NaN)
    """

    # --- ensure numeric types so == works ---------------------------------
    for col in ("label_id", "column_id"):
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    commits = df[
        ((df["label_id"] == 0) & (df["column_id"] == 1)) |
        (df.get("board_name", "") == "COMPROMISOS")
    ].copy()

    if commits.empty:
        doc.add_paragraph("No commitments recorded.")
        return

    # --- Word table header -------------------------------------------------
    table = doc.add_table(rows=1, cols=3)
    table.style  = "Table Grid"
    table.autofit = False
    hdr_row = table.rows[0]
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr_row.height      = Inches(0.45)
    table.columns[0].width = Inches(3.0)   # COMPROMISO
    table.columns[1].width = Inches(4.0)   # RESPONSABLE
    table.columns[2].width = Inches(3.0)   # FECHA

    for i, text in enumerate(("COMPROMISO", "RESPONSABLE", "FECHA")):
        cell = hdr_row.cells[i]
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold, run.font.size, run.font.name = True, Pt(12), "Verdana"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), BRAND_COLOR_HEADER)
        cell._element.get_or_add_tcPr().append(shd)

    # --- data rows ---------------------------------------------------------
    for _, row in commits.iterrows():
        new_cells = table.add_row().cells
    
        if row.get("board_name") == "COMPROMISOS":
            # ── legacy mapping ───────────────────────────────────────────
            comp  = str(row.get("title", ""))                 # Compromiso
            resp  = str(row.get("planlet_name", ""))          # Responsable
            raw_c = str(row.get("comments_parsed", ""))       # Fecha
            fecha = parse_comment_for_date(raw_c) or raw_c.strip("[]'\" ") or "N/A"
    
        elif pd.notna(row.get("label_id")) and int(row.get("label_id")) == 0:
            # ── new mapping (label_id == 0) ──────────────────────────────
            comp  = str(row.get("title", ""))                  # COMPROMISO
            resp  = str(row.get("comments_parsed", ""))        # RESPONSABLE
            fecha = safe_parse_due(row.get("due_date"))        # FECHA
    
        else:
            # Row doesn’t match either style → remove the extra row & skip
            table._tbl.remove(new_cells[0]._tc)
            continue
    
        # write the three values into the row
        for cell, value in zip(new_cells, (comp, resp, fecha)):
            cell.text = value
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.runs[0]
            run.font.size, run.font.name = Pt(10), "Verdana"

            
def parse_comment_for_date(comment_text):
    c = comment_text.strip("[]'\" ")
    for fmt in ("%d/%m/%Y", "%m/%d/%Y"):
        try:
            dt = datetime.strptime(c, fmt)
            return str(dt.date())
        except ValueError:
            pass
    return ""

# ----------------------------------------------------------------------------
# S3 & HELPER FUNCS
# ----------------------------------------------------------------------------
def upload_file_to_s3(file_path, s3_key):
    """
    Upload with ContentType and ContentDisposition so that direct S3
    console downloads preserve the docx or xlsx properly.
    """
    s3 = boto3.client("s3", region_name=REGION)
    content_type = infer_content_type(s3_key)
    disposition = f'attachment; filename="{os.path.basename(s3_key)}"'

    try:
        s3.upload_file(
            Filename=file_path,
            Bucket=S3_BUCKET,
            Key=s3_key,
            ExtraArgs={
                "ContentType": content_type,
                "ContentDisposition": disposition
            }
        )
        logger.info(f"✅ Uploaded {file_path} => s3://{S3_BUCKET}/{s3_key}")
        return True
    except Exception as e:
        logger.error(f"❌ S3 upload error: {str(e)}")
        return False


def convert_docx_to_pdf(doc_path):
    """
    Uses LibreOffice in headless mode to convert a .docx -> .pdf
    E.g. doc_path=/tmp/Acta_Project123.docx
    Output => /tmp/Acta_Project123.pdf
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Docx not found: {doc_path}")

    output_dir = os.path.dirname(doc_path)
    # LibreOffice command => `libreoffice --headless --convert-to pdf input.docx --outdir /tmp`
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        doc_path,
        "--outdir",
        output_dir
    ]
    subprocess.run(cmd, check=True)  # Raises CalledProcessError if fail

    pdf_path = doc_path.replace(".docx", ".pdf")
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF not found after conversion: {pdf_path}")

    logger.info(f"PDF created => {pdf_path}")
    return pdf_path


def infer_content_type(s3_key):
    if s3_key.lower().endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif s3_key.lower().endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif s3_key.lower().endswith(".pdf"):
        return "application/pdf"
    else:
        return "application/octet-stream"

def set_document_margins_and_orientation(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

def add_page_x_of_y_footer(doc):
    section = doc.sections[0]
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Start: "Page "
    run = footer_para.add_run("Page ")
    add_field(run, "PAGE")

    # " of "
    run = footer_para.add_run(" of ")
    add_field(run, "NUMPAGES")

def add_field(run, field_code):
    """
    Insert a Word field, e.g. 'PAGE' or 'NUMPAGES', into a run.
    """
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    # Mark field as "dirty" so Word updates it on open
    dirty = OxmlElement('w:fldChar')
    dirty.set(qn('w:fldCharType'), 'end')

    # Append all
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(dirty)


def add_top_header_table(doc, main_title, logo_path=None):
    """
    2 columns => each 5.0"
    Title => left col, optional logo => right col
    """
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(5.0)

    left_cell = table.cell(0,0)
    p_left = left_cell.paragraphs[0]
    run_left = p_left.add_run(main_title)
    run_left.bold = True
    run_left.font.size = Pt(30)
    run_left.font.name = "Verdana"

    right_cell = table.cell(0,1)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    if logo_path and os.path.exists(logo_path):
        run_logo = p_right.add_run()
        run_logo.add_picture(logo_path, width=Inches(3.0))
    else:
        fallback_run = p_right.add_run("[INSERT COMPANY LOGO HERE]")
        fallback_run.bold = True
        fallback_run.font.size = Pt(12)


def add_two_by_two_table(doc, date_text, project_name, project_id, leader_name, shade_color=None):
    """
    2x2 => (FECHA DEL ACTA, PROYECTO) / (NO. PROYECTO, PROJECT MANAGER)
    """
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(5.0)

    # row0 col0 => FECHA DEL ACTA
    c_0_0 = table.cell(0,0)
    r_0_0 = c_0_0.paragraphs[0].add_run(date_text)
    r_0_0.font.size = Pt(12)
    r_0_0.bold = True
    if shade_color:
        shade_cell(c_0_0, shade_color)

    # row0 col1 => PROYECTO
    c_0_1 = table.cell(0,1)
    r_0_1 = c_0_1.paragraphs[0].add_run(f"PROYECTO  {project_name}")
    r_0_1.font.size = Pt(12)
    r_0_1.bold = True
    if shade_color:
        shade_cell(c_0_1, shade_color)

    # row1 col0 => NO. PROYECTO
    c_1_0 = table.cell(1,0)
    r_1_0 = c_1_0.paragraphs[0].add_run(f"NO. DE PROYECTO  {project_id}")
    r_1_0.font.size = Pt(12)
    r_1_0.bold = True
    if shade_color:
        shade_cell(c_1_0, shade_color)

    # row1 col1 => PROJECT MANAGER
    c_1_1 = table.cell(1,1)
    r_1_1 = c_1_1.paragraphs[0].add_run(f"PROJECT MANAGER  {leader_name}")
    r_1_1.font.size = Pt(12)
    r_1_1.bold = True
    if shade_color:
        shade_cell(c_1_1, shade_color)


def add_asistencia_table(doc, df):
    """
    2x2 => row0 col0 => "ASISTENCIA", row0 col1 => "ASISTENCIA CLIENTE"
            row1 col0 => planlet_name="ASISTENCIA"
            row1 col1 => planlet_name="ASISTENCIA CLIENTE"
    """
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(5.0)

    # row0 => Headers
    hdr_row = table.rows[0]
    hdr_cells = hdr_row.cells

    # CHANGED: "ASISTENCIA" → "ASISTENCIA CLIENTE"
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[0].text = "ASISTENCIA IKUSI"  # <--- CHANGED
    run0 = hdr_cells[0].paragraphs[0].runs[0]
    run0.bold = True
    run0.font.size = Pt(14)
    run0.font.name = "Verdana"
    run0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shading_elm0 = OxmlElement("w:shd")
    shading_elm0.set(qn("w:fill"), BRAND_COLOR_HEADER)
    hdr_cells[0]._element.get_or_add_tcPr().append(shading_elm0)

    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].text = "ASISTENCIA CLIENTE"
    run1 = hdr_cells[1].paragraphs[0].runs[0]
    run1.bold = True
    run1.font.size = Pt(14)
    run1.font.name = "Verdana"
    run1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shading_elm1 = OxmlElement("w:shd")
    shading_elm1.set(qn("w:fill"), BRAND_COLOR_HEADER)
    hdr_cells[1]._element.get_or_add_tcPr().append(shading_elm1)

    # row1 => data from planlet_name= "ASISTENCIA" & planlet_name= "ASISTENCIA CLIENTE"
    row_asist = df[(df["planlet_name"]=="ASISTENCIA") & (df["column_id"]==1)]
    text_asist = ""
    if not row_asist.empty:
        text_asist = str(row_asist.iloc[0].get("comments_parsed",""))

    row_asist_cliente = df[(df["planlet_name"]=="ASISTENCIA CLIENTE") & (df["column_id"]==1)]
    text_asist_cliente = ""
    if not row_asist_cliente.empty:
        text_asist_cliente = str(row_asist_cliente.iloc[0].get("comments_parsed",""))

    data_row = table.rows[1].cells
    data_row[0].text = text_asist
    data_row[1].text = text_asist_cliente

    # style
    for col_idx in range(2):
        p = data_row[col_idx].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_data = p.runs[0]
        run_data.font.size = Pt(10)
        run_data.font.name = "Verdana"


def add_section_header(doc, header_text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_h = p.add_run(header_text)
    run_h.bold = True
    run_h.font.size = Pt(18)
    run_h.font.name = "Verdana"
    doc.add_paragraph()


def add_horizontal_rule(doc):
    line_para = doc.add_paragraph()
    line_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = line_para.add_run("_______________________________________________________________")
    run.font.size = Pt(12)


def shade_cell(cell, color):
    shade_elm = OxmlElement("w:shd")
    shade_elm.set(qn("w:fill"), color)
    cell._element.get_or_add_tcPr().append(shade_elm)

def add_unified_visual_header(doc, main_title, logo_path=None):
    """
    Final design per visual spec:
    - Merge cell(0,0) + cell(1,0) for the logo (vertical span)
    - Merge cell(0,1) + cell(0,2) for centered title
    - Row 2: metadata fields left-to-right
    """
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    table.columns[0].width = Inches(2.5)  # Logo
    table.columns[1].width = Inches(2.5)  # Title
    table.columns[2].width = Inches(2.0)  # Spacer
    table.columns[3].width = Inches(3.0)  # Code block

    # Merge for logo vertical (row 0 & 1 col 0)
    logo_cell = table.cell(0, 0).merge(table.cell(1, 0))
    if logo_path and os.path.exists(logo_path):
        run_logo = logo_cell.paragraphs[0].add_run()
        run_logo.add_picture(logo_path, width=Inches(2.0))
    else:
        logo_cell.text = "LOGO"

    # Merge for ACTA title (row 0 col 1+2)
    title_cell = table.cell(0, 1).merge(table.cell(0, 2))
    para_title = title_cell.paragraphs[0]
    para_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = para_title.add_run(main_title)
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.name = "Verdana"

    # Código and Fecha cell (row 0 col 3)
    code_para = table.cell(0, 3).paragraphs[0]
    run_code = code_para.add_run("Código: GP-F-004\nFecha: 13-02-2020")
    run_code.font.size = Pt(10)
    run_code.font.name = "Verdana"

    # Row 2 metadata
    table.cell(1, 1).text = "Revisó: Gerente de Operaciones"
    table.cell(1, 2).text = "Aprobó: Gestión Documental"
    table.cell(1, 3).text = "Versión: 2"

    # Apply formatting
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in para.runs:
                    run.font.name = "Verdana"

    doc.add_paragraph()
