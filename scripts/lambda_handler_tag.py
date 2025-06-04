# lambda_handler_tag.py – full end‑to‑end script with new tag‑based COMPROMISOS logic
# ---------------------------------------------------------------------------
#  ▓▓  KEY CHANGES vs current prod
#  1. Fetch tag catalogue per project (endpoint /1/tags/projects/{pid}/cards)
#  2. Identify cards whose tag name == "compromiso" (case‑insensitive)
#  3. Keep card.due_date  → Fecha column
#  4. COMPROMISOS table rows built from df[df["is_compromiso"]]
#  5. No dependency on board == "COMPROMISOS" ⚠️  (backwards compatible)
# ---------------------------------------------------------------------------

import os, json, time, logging, requests, ast, subprocess
from datetime import datetime
import pandas as pd

import boto3
from botocore.exceptions import ClientError

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --------------------------- CONFIG ---------------------------
SECRET_NAME          = "ProjectPlaceAPICredentials"
REGION               = "us-east-2"
API_ROOT             = "https://api.projectplace.com"
OUTPUT_EXCEL         = "/tmp/Acta_de_Seguimiento.xlsx"
DYNAMO_TABLE         = os.getenv("DYNAMODB_TABLE_NAME", "ProjectPlace_DataExtrator_landing_table_v3")
S3_BUCKET            = os.getenv("S3_BUCKET_NAME", "projectplace-dv-2025-x9a7b")
BRAND_COLOR_HEADER   = "4AC795"
LIGHT_SHADE_2X2      = "FAFAFA"
LOGO_IMAGE_PATH      = os.path.join(os.path.dirname(__file__), "logo", "company_logo.png")

# --------------------------- LAMBDA ENTRY ---------------------

def lambda_handler(event, context):
    logger.info("Lambda start – tag‑aware Acta build")

    # 1. Secrets
    sec = _load_secrets()
    tok = _oauth(sec.get("PROJECTPLACE_ROBOT_CLIENT_ID"), sec.get("PROJECTPLACE_ROBOT_CLIENT_SECRET"))
    if not tok:
        return {"statusCode":500, "body":"Token error"}

    # 2. Projects
    projects = _projects(tok)
    if not projects:
        return {"statusCode":200, "body":"No active projects"}

    # 3. Excel generation
    excel = _excel(tok, projects)
    if not excel:
        return {"statusCode":500, "body":"Excel failed"}

    df = pd.read_excel(excel)
    logger.info(f"Excel → rows={len(df)} cols={list(df.columns)}")
    if df.empty:
        return {"statusCode":200, "body":"Empty Excel"}

    _ddb_persist(df)
    df = _filter(df)
    logger.info(f"After filter rows={len(df)}")
    if df.empty:
        return {"statusCode":200, "body":"Nothing after filter"}

    # 4. DOCs per project
    made=0
    for pid, p_df in df.groupby("project_id"):
        logger.info(f"Build acta pid={pid} rows={len(p_df)}")
        path = _build_doc(pid, p_df)
        if not path:
            logger.error(f"build_doc returned None for {pid}"); continue
        made += 1
        safe = str(p_df.iloc[0]["project_name"]).replace(" ","_").replace("/","_")
        _s3_upload(path, f"actas/Acta_{safe}_{pid}.docx")

    _s3_upload(excel, "actas/Acta_de_Seguimiento.xlsx")
    return {"statusCode":200, "body":json.dumps({"docs":made})}

# --------------------------- HELPERS -------------------------

def _load_secrets():
    sm=boto3.client("secretsmanager",region_name=REGION)
    try:return json.loads(sm.get_secret_value(SecretId=SECRET_NAME)["SecretString"])
    except ClientError as e:
        logger.error(e);return {}

def _oauth(cid,sec):
    if not cid or not sec: return None
    r=requests.post(f"{API_ROOT}/oauth2/access_token",data={"grant_type":"client_credentials","client_id":cid,"client_secret":sec})
    try: r.raise_for_status(); return r.json().get("access_token")
    except Exception as e: logger.error(e); return None

def _projects(tok):
    h={"Authorization":f"Bearer {tok}"}
    r=requests.get(f"{API_ROOT}/1/account/projects",headers=h); r.raise_for_status()
    return [p for p in r.json().get("projects",[]) if not p.get("archived",False)]

def _tag_lookup(tok,pid):
    """Return set of tag IDs whose name == 'compromiso' (lower)."""
    h={"Authorization":f"Bearer {tok}"}
    r=requests.get(f"{API_ROOT}/1/tags/projects/{pid}/cards",headers=h)
    r.raise_for_status()
    return {t["id"] for t in r.json() if t["name"].lower()=="compromiso"}

def _card_comments(tok,cid):
    h={"Authorization":f"Bearer {tok}"}
    r=requests.get(f"{API_ROOT}/1/cards/{cid}/comments",headers=h)
    r.raise_for_status(); return [c.get("text","") for c in r.json()]

# --------------------------- EXCEL ---------------------------

def _excel(tok, projects):
    rows=[]; h={"Authorization":f"Bearer {tok}"}
    for proj in projects:
        pid,pname=str(proj["id"]),proj.get("name","Unnamed")
        compromiso_ids = _tag_lookup(tok,pid)
        r=requests.get(f"{API_ROOT}/1/projects/{pid}/cards",headers=h); r.raise_for_status()
        for c in r.json():
            row=dict(c)
            row.update({
                "due_date"     : c.get("due_date"),
                "is_compromiso": any(t in compromiso_ids for t in c.get("tag_ids",[])),
                "Comments"     : str(_card_comments(tok,c["id"])),
                "project_id"   : pid,
                "project_name" : pname
            })
            rows.append(row)
    if not rows: return None
    pd.DataFrame(rows).to_excel(OUTPUT_EXCEL,index=False)
    return OUTPUT_EXCEL

# --------------------------- DDB ----------------------------

def _ddb_persist(df):
    tbl=boto3.resource("dynamodb",region_name=REGION).Table(DYNAMO_TABLE)
    for _,r in df.iterrows():
        tbl.put_item(Item={
            "project_id":str(r["project_id"]),
            "card_id"   :str(r["id"]),
            "title"     :str(r.get("title","")),
            "due_date"  :str(r.get("due_date","")),
            "timestamp" :int(time.time())
        })

# --------------------------- FILTER --------------------------

def _filter(df):
    if "column_id" in df.columns:
        df=df[df["column_id"]==1]
    df["comments_parsed"] = df["Comments"].apply(lambda raw:(ast.literal_eval(raw)[-1] if raw.startswith("[") else raw))
    return df

# -------------------- DOC (python‑docx) ----------------------

def _build_doc(pid, pdf):
    doc=Document(); _layout(doc); _header(doc)
    _commitments_table(doc,pdf)
    path=f"/tmp/Acta_{pid}.docx"; doc.save(path); return path

# layout, header helpers kept minimal to focus on COMPROMISOS ----------------

def _layout(doc):
    s=doc.sections[0]; s.orientation=WD_ORIENT.LANDSCAPE; s.page_width,s.page_height=Inches(11),Inches(8.5)
    for m in ("left_margin","right_margin","top_margin","bottom_margin"):setattr(s,m,Inches(0.5))

def _header(doc):
    t=doc.add_table(rows=2,cols=4); t.style="Table Grid"; t.autofit=False
    t.columns[0].width,t.columns[1].width,t.columns[2].width,t.columns[3].width=[Inches(2.5),Inches(2.5),Inches(2),Inches(3)]
    logo=t.cell(0,0).merge(t.cell(1,0))
    if os.path.exists(LOGO_IMAGE_PATH): logo.paragraphs[0].add_run().add_picture(LOGO_IMAGE_PATH,width=Inches(2))
    else: logo.text="LOGO"
    title=t.cell(0,1).merge(t.cell(0,2)); p=title.paragraphs[0]; p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    r=p.add_run("ACTA DE SEGUIMIENTO"); r.bold=True; r.font.size=Pt(20); r.font.name="Verdana"
    t.cell(0,3).text="Código: GP-F-004\nFecha: 13-02-2020"
    t.cell(1,1).text="Revisó: Gerente de Operaciones"; t.cell(1,2).text="Aprobó: Gestión Documental"; t.cell(1,3).text="Versión: 2"

# -------------------- COMPROMISOS TABLE ---------------------

def _commitments_table(doc, df):
    df=df[df["is_compromiso"]]
    if df.empty:
        doc.add_paragraph("No commitments recorded."); return
    tbl=doc.add_table(rows=1,cols=3); tbl.style="Table Grid"; tbl.autofit=False
    tbl.columns[0].width,tbl.columns[1].width,tbl.columns[2].width=[Inches(3),Inches(4),Inches(3)]
    hdrs=["COMPROMISO","RESPONSABLE","FECHA"]
    for i,h in enumerate(hdrs):
        c=tbl.rows[0].cells[i]; c.text=h; r=c.paragraphs[0].runs[0]; r.bold,r.font.size,r.font.name=True,Pt(12),"Verdana"
        sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),BRAND_COLOR_HEADER); c._element.get_or_add_tcPr().append(sh)
    for _,row in df.iterrows():
        comp=row.get("title","")
        resp=str(row.get("comments_parsed","")).strip("[]'\"")
        raw=row.get("due_date","")
        try: fecha=datetime.strptime(raw[:10],"%Y-%m-%d").strftime("%d/%m/%Y")
        except Exception: fecha=raw
        for val,cell in zip([comp,resp,fecha], tbl.add_row().cells):
            cell.text=val; cell.paragraphs[0].runs[0].font.size,cell.paragraphs[0].runs[0].font.name=Pt(10),"Verdana"

# --------------------------- S3 -----------------------------

def _s3_upload(path,key):
    s3=boto3.client("s3",region_name=REGION)
    s3.upload_file(path,S3_BUCKET,key,ExtraArgs={"ContentType":_ctype(key),"ContentDisposition":f"attachment; filename={os.path.basename(key)}"})
    logger.info(f"Uploaded {key}")

def _ctype(key):
    if key.endswith(".docx"): return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if key.endswith(".xlsx"): return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "application/octet-stream"
